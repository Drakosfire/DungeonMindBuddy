from __future__ import annotations

import re
from pathlib import Path

import pytest
from docx import Document

from src.ingestion.docx_converter import docx_to_markdown, markdown_passthrough


def test_docx_to_markdown_style_and_fallback_detection(tmp_path: Path) -> None:
    doc = Document()
    doc.add_heading("Mirathorn Overview", level=1)
    doc.add_paragraph("Founded over 200 years ago.")
    doc.add_paragraph("1. Watch Tower")
    p_bold = doc.add_paragraph()
    p_bold.add_run("Temple of the Nameless Stone").bold = True
    p_text = doc.add_paragraph()
    p_text.add_run("A ")
    p_text.add_run("prosperous").bold = True
    p_text.add_run(" city.")

    out_path = tmp_path / "sample.docx"
    doc.save(out_path)

    markdown = docx_to_markdown(out_path)

    assert "# Mirathorn Overview" in markdown
    assert "## 1. Watch Tower" in markdown
    assert "## Temple of the Nameless Stone" in markdown
    assert "A **prosperous** city." in markdown


def test_markdown_passthrough_reads_file(tmp_path: Path) -> None:
    md_path = tmp_path / "sample.md"
    md_path.write_text("# Title\n\nBody\n", encoding="utf-8")
    assert markdown_passthrough(md_path) == "# Title\n\nBody\n"


def test_mirathorn_docx_has_detectable_headings() -> None:
    mirathorn_path = Path(
        "/media/drakosfire/Projects/DungeonOverMind/DungeonMindBuddy/Docs/"
        "Eldyrwild and Campaign Context/Elderwyld/Cities and Towns/Mirathorn/"
        "The City of Mirathorn.docx"
    )
    if not mirathorn_path.exists():
        pytest.skip("Mirathorn corpus docx is not present on this machine")

    markdown = docx_to_markdown(mirathorn_path)
    heading_count = len(re.findall(r"(?m)^#{1,4}\s+.+$", markdown))
    assert heading_count >= 1
